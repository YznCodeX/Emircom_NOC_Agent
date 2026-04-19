"""
reports.py — Word (.docx) and Excel (.xlsx) report generators
=============================================================

What this file is responsible for
----------------------------------
Everything needed to produce the two downloadable shift-handoff reports
that appear in Tab 2 → "📄 Handoff Report" section of the dashboard.
All heavy document-building logic was extracted here so that app.py
stays focused on UI wiring.

Rules for this module
---------------------
• No Streamlit imports — functions return BytesIO objects that the
  caller passes straight to st.download_button(); nothing is rendered here.
• Depends on: python-docx, openpyxl, pandas, the standard LLM shim
  (src.agent_graph.llm) for the AI-written summary section.
• Called by: app.py (Tab 2 download buttons).

Functions
---------
  _set_cell_bg(cell, hex_colour: str) -> None   [private]
      Applies a background fill to a Word table cell using raw XML.
      Used to colour the header rows of tables in the report.

  _styled_heading(doc, text: str, level: int, colour: RGBColor) -> None   [private]
      Adds a heading paragraph with a custom RGB colour instead of the
      default Word theme colour.

  _add_table(doc, headers, rows, header_bg) -> None   [private]
      Helper that builds a styled Word table with a coloured header row
      and alternating row shading (light grey / white).

  run_handoff_llm(processed: list[dict]) -> str
      Calls the Groq LLM to write a 3-paragraph shift summary in the
      style of a senior NOC engineer.  Returns a plain-text string.
      Falls back to a short error message if the LLM call fails.

  generate_handoff_report_doc(processed: list[dict]) -> BytesIO
      Builds the full Word document:
        • Cover section (shift metadata, engineer name placeholder)
        • AI-generated shift summary (from run_handoff_llm)
        • Ticket summary table (all processed tickets)
        • SLA compliance table
        • Recommendations section
      Returns the document as an in-memory BytesIO ready for download.

  generate_excel_report(processed: list[dict]) -> BytesIO
      Writes the processed ticket list to an Excel workbook with a
      header row and auto-fitted column widths.
      Returns a BytesIO ready for st.download_button().
"""
import io
import json
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Word helpers ───────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Apply a solid background fill colour to a Word table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tc_pr.append(shd)


def _styled_heading(doc, text: str, level: int = 1):
    """Add a heading with consistent Emircom dark-blue colour."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def _add_table(doc, headers: list, rows: list, col_widths: list | None = None):
    """Dark-blue header row + alternating white/light-blue data rows."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[i], "1F4E79")

    light_fill = "EBF3FB"
    for idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row_data):
            cells[j].text = str(val)
            for para in cells[j].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(9)
            if idx % 2 == 1:
                _set_cell_bg(cells[j], light_fill)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


# ── LLM narrative generator ────────────────────────────────────────────────────

def run_handoff_llm(tickets_data: list, pending_tickets: list) -> dict:
    """
    Single LLM call that writes the three narrative sections of the handoff report:
    - shift_narrative
    - critical_summaries  (dict keyed by ticket ID)
    - watch_list
    """
    from src.agent_graph import llm

    processed_summary = json.dumps(tickets_data, indent=2)
    pending_summary   = json.dumps([
        {"Ticket_ID": t["Ticket_ID"], "Category": t["Category"],
         "Alert_Message": t.get("Alert_Message", "")}
        for t in pending_tickets
    ], indent=2)
    critical_ids = [t["Ticket_ID"] for t in tickets_data
                    if str(t.get("Severity", "")).upper() == "CRITICAL"]

    prompt = f"""You are a senior NOC Shift Lead at Emircom writing a formal end-of-shift handover report.

PROCESSED TICKETS THIS SHIFT:
{processed_summary}

PENDING (NOT YET PROCESSED) TICKETS:
{pending_summary}

CRITICAL TICKET IDs: {critical_ids}

Generate three sections. Reply ONLY with valid JSON — no markdown, no explanation:

{{
  "shift_narrative": "3-5 sentences. Cover: overall ticket volume, dominant categories, major patterns or correlations observed, SLA performance, and an overall severity assessment of the shift. Technical, direct, no filler.",
  "critical_summaries": {{
    "INC-XXXX": "2-3 sentences per critical ticket: what happened, what action was taken, current resolution status. Reference actual ticket data. If there are no critical tickets return an empty object {{}}"
  }},
  "watch_list": "3-5 sentences. Tell the incoming engineer specifically what to monitor: which devices, services, or alert patterns from this shift may re-occur or escalate. Reference actual ticket IDs, device names, or category trends from the data."
}}"""

    response = llm.invoke(prompt)
    raw = response.content.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)


# ── Word document builder ──────────────────────────────────────────────────────

def generate_handoff_report_doc(
    tickets_data:    list,
    pending_tickets: list,
    outgoing_eng:    str,
    incoming_eng:    str,
    shift_period:    str,
    ai_content:      dict,
    short_mode:      bool = False,
) -> io.BytesIO:
    """Build the full NOC Shift Handoff .docx and return it as a BytesIO buffer."""

    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Inches(0.9)
        sec.bottom_margin = Inches(0.9)
        sec.left_margin   = Inches(1.1)
        sec.right_margin  = Inches(1.1)

    sev_colors = {"CRITICAL": "C00000", "HIGH": "C55A11", "MEDIUM": "7F6000", "LOW": "375623"}
    sev_fills  = {"CRITICAL": "FF0000", "HIGH": "FF8C00", "MEDIUM": "FFC000", "LOW": "70AD47"}

    # ── Cover block ────────────────────────────────────────────────────────────
    def _cp(text, size=11, bold=False, italic=False,
            align=WD_ALIGN_PARAGRAPH.CENTER, rgb: RGBColor | None = None):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.bold      = bold
        run.italic    = italic
        run.font.size = Pt(size)
        if rgb:
            run.font.color.rgb = rgb
        return p

    _cp("EMIRCOM NOC / SOC COMMAND CENTER", size=18, bold=True,  rgb=RGBColor(0x1F, 0x4E, 0x79))
    _cp("SHIFT HANDOFF REPORT",             size=13, bold=True,  rgb=RGBColor(0x2E, 0x74, 0xB5))
    _cp("Confidential — Internal Use Only", size=9,  italic=True,rgb=RGBColor(0x80, 0x80, 0x80))
    doc.add_paragraph()

    info_table = doc.add_table(rows=5, cols=2)
    info_data  = [
        ("Shift Date",        datetime.now().strftime("%d %B %Y")),
        ("Shift Period",      shift_period or "—"),
        ("Outgoing Engineer", outgoing_eng or "—"),
        ("Incoming Engineer", incoming_eng or "—"),
        ("Report Generated",  datetime.now().strftime("%H:%M")),
    ]
    for i, (label, value) in enumerate(info_data):
        lbl_cell, val_cell = info_table.rows[i].cells[0], info_table.rows[i].cells[1]
        lbl_cell.text = label
        val_cell.text = value
        lbl_run = lbl_cell.paragraphs[0].runs[0]
        lbl_run.bold            = True
        lbl_run.font.size       = Pt(10)
        lbl_run.font.color.rgb  = RGBColor(0x1F, 0x4E, 0x79)
        val_cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ── Section 1: Shift Narrative (full mode only) ────────────────────────────
    if not short_mode:
        _styled_heading(doc, "1.  Shift Narrative")
        narrative = ai_content.get("shift_narrative", "No narrative generated.")
        p = doc.add_paragraph(narrative)
        p.runs[0].font.size = Pt(10.5)
        doc.add_paragraph()

    # ── Section 2 / 1: Key Metrics ────────────────────────────────────────────
    _styled_heading(doc, "1.  Key Metrics" if short_mode else "2.  Key Metrics")
    total        = len(tickets_data)
    approved     = sum(1 for t in tickets_data if t["Status"] in ("Approved", "Approved (Queue)"))
    dropped      = sum(1 for t in tickets_data if t["Status"] == "Dropped (Duplicate)")
    rejected     = sum(1 for t in tickets_data if t["Status"] == "Rejected")
    sla_breached = sum(1 for t in tickets_data if t.get("SLA_Breached", False))
    crit_cnt     = sum(1 for t in tickets_data if str(t.get("Severity","")).upper() == "CRITICAL")
    high_cnt     = sum(1 for t in tickets_data if str(t.get("Severity","")).upper() == "HIGH")
    pending_cnt  = len(pending_tickets)

    # Derived metrics
    rt_values      = [int(t.get("Response_Time_Secs", 0) or 0) for t in tickets_data]
    avg_rt_secs    = int(sum(rt_values) / len(rt_values)) if rt_values else 0
    avg_rt_str     = f"{avg_rt_secs // 60}m {avg_rt_secs % 60}s" if avg_rt_secs else "—"
    compliance_pct = f"{round((1 - sla_breached / total) * 100)}%" if total else "—"
    conf_values    = [float(t["Confidence_Score"]) for t in tickets_data
                      if str(t.get("Confidence_Score", "")).strip().lstrip("-").isdigit()]
    avg_conf       = f"{round(sum(conf_values) / len(conf_values))}%" if conf_values else "—"

    _add_table(doc,
        headers=["Metric", "Value"],
        rows=[
            ("Total Tickets Processed",  total),
            ("Approved & Escalated",     approved),
            ("Duplicates Dropped",       dropped),
            ("Rejected",                 rejected),
            ("SLA Breaches",             sla_breached),
            ("SLA Compliance Rate",      compliance_pct),
            ("Critical Incidents",       crit_cnt),
            ("High-Severity Incidents",  high_cnt),
            ("Avg. Response Time",       avg_rt_str),
            ("Avg. AI Confidence Score", avg_conf),
            ("Pending (Not Processed)",  pending_cnt),
        ],
        col_widths=[3.0, 1.0],
    )
    doc.add_paragraph()

    # ── Severity Breakdown ────────────────────────────────────────────────────
    _styled_heading(doc, "Severity Breakdown", level=2)
    sev_levels = ["CRITICAL", "HIGH", "MEDIUM", "LOW"]
    sev_rows = []
    for s in sev_levels:
        cnt = sum(1 for t in tickets_data if str(t.get("Severity","")).upper() == s)
        pct = f"{round(cnt / total * 100)}%" if total else "0%"
        sev_rows.append((s.title(), cnt, pct))
    _add_table(doc,
        headers=["Severity", "Count", "% of Total"],
        rows=sev_rows,
        col_widths=[1.6, 0.8, 1.2],
    )
    doc.add_paragraph()

    # ── Section 3 / 2: Open Items ─────────────────────────────────────────────
    _styled_heading(doc, "2.  Open Items — Requires Incoming Engineer Attention"
                         if short_mode else "3.  Open Items — Requires Incoming Engineer Attention")

    if not pending_tickets:
        p = doc.add_paragraph("✅  Queue is clear. No open items for the incoming shift.")
        p.runs[0].font.size       = Pt(10.5)
        p.runs[0].font.color.rgb  = RGBColor(0x37, 0x56, 0x23)
    else:
        warn = doc.add_paragraph(
            f"⚠️  {len(pending_tickets)} ticket(s) were NOT processed this shift "
            "and must be handled immediately by the incoming engineer."
        )
        warn.runs[0].bold           = True
        warn.runs[0].font.size      = Pt(10.5)
        warn.runs[0].font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        doc.add_paragraph()
        _add_table(doc,
            headers=["Ticket ID", "Category", "Alert / Description"],
            rows=[(t["Ticket_ID"], t["Category"], t.get("Alert_Message", "")[:80])
                  for t in pending_tickets],
            col_widths=[1.2, 1.2, 4.6],
        )
    doc.add_paragraph()

    # ── Sections 4-6: Full mode only ──────────────────────────────────────────
    if not short_mode:

        # Section 4: Critical Incident Summaries
        _styled_heading(doc, "4.  Critical Incident Summaries")
        critical_summaries: dict = ai_content.get("critical_summaries", {})
        critical_tickets = [t for t in tickets_data if str(t.get("Severity","")).upper() == "CRITICAL"]

        if not critical_tickets:
            p = doc.add_paragraph("No Critical incidents recorded this shift.")
            p.runs[0].font.size = Pt(10.5)
        else:
            for t in critical_tickets:
                tid = t.get("Ticket_ID", "")
                h2  = doc.add_heading("", level=2)
                run = h2.add_run(f"{tid}  —  {t.get('Category','')}  |  Status: {t.get('Status','')}")
                run.bold           = True
                run.font.size      = Pt(11)
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                p = doc.add_paragraph(critical_summaries.get(tid, "No AI summary available."))
                p.runs[0].font.size = Pt(10.5)
                if t.get("Correlated_With"):
                    cn = doc.add_paragraph(f"↳  Correlated with: {t['Correlated_With']}")
                    cn.runs[0].font.size      = Pt(9.5)
                    cn.runs[0].italic         = True
                    cn.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x80)
                doc.add_paragraph()

        # Section 5: Correlated Incident Groups
        _styled_heading(doc, "5.  Correlated Incident Groups")
        corr_groups: dict = {}
        for t in tickets_data:
            cw = t.get("Correlated_With", "")
            if cw:
                corr_groups.setdefault(cw, set()).add(cw)
                corr_groups[cw].add(t["Ticket_ID"])

        if not corr_groups:
            p = doc.add_paragraph("No correlated incident groups detected this shift.")
            p.runs[0].font.size = Pt(10.5)
        else:
            for group_label, (root, members) in enumerate(corr_groups.items(), start=1):
                gh = doc.add_paragraph()
                gr = gh.add_run(f"Group {group_label}  —  Root: {root}")
                gr.bold            = True
                gr.font.size       = Pt(10.5)
                gr.font.color.rgb  = RGBColor(0x1F, 0x4E, 0x79)
                mp = doc.add_paragraph(f"Linked tickets: {',  '.join(sorted(members))}")
                mp.runs[0].font.size = Pt(10)
                mp.runs[0].italic    = True
                doc.add_paragraph()

        # Section 6: Full Incident Log
        _styled_heading(doc, "6.  Full Incident Log")
        log_headers = ["Ticket ID", "Category", "Severity", "Status",
                       "Response Time", "SLA Breached", "Correlated With", "Confidence"]
        log_table = doc.add_table(rows=1, cols=len(log_headers))
        log_table.style = "Light Grid Accent 1"

        hdr_cells = log_table.rows[0].cells
        for i, h in enumerate(log_headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold           = True
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(hdr_cells[i], "1F4E79")

        light = "EBF3FB"
        for idx, t in enumerate(tickets_data):
            rt_secs  = int(t.get("Response_Time_Secs", 0) or 0)
            rt_str   = f"{rt_secs // 60}m {rt_secs % 60}s" if rt_secs else "—"
            severity = str(t.get("Severity", "")).upper()
            cells    = log_table.add_row().cells
            cells[0].text = str(t.get("Ticket_ID",      ""))
            cells[1].text = str(t.get("Category",        ""))
            cells[2].text = severity
            cells[3].text = str(t.get("Status",          ""))
            cells[4].text = rt_str
            cells[5].text = "YES" if t.get("SLA_Breached", False) else "No"
            cells[6].text = str(t.get("Correlated_With", "") or "—")
            cells[7].text = str(t.get("Confidence_Score","") or "—")

            fill = sev_fills.get(severity, "CCCCCC")
            _set_cell_bg(cells[2], fill)
            runs = cells[2].paragraphs[0].runs
            if runs:
                runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                runs[0].bold = True
            if idx % 2 == 1:
                for j, cell in enumerate(cells):
                    if j != 2:
                        _set_cell_bg(cell, light)
            for cell in cells:
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.font.size = Pt(9)

        doc.add_paragraph()

    # ── Watch List (both modes) ────────────────────────────────────────────────
    _styled_heading(doc, "3.  Watch List — Incoming Shift Priorities"
                         if short_mode else "7.  Watch List — Incoming Shift Priorities")
    p = doc.add_paragraph(ai_content.get("watch_list", "No watch list generated."))
    p.runs[0].font.size = Pt(10.5)
    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────────────
    sep = doc.add_paragraph("─" * 72)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot = doc.add_paragraph(
        f"Generated by Emircom NOC AI Agent  •  {datetime.now().strftime('%d %B %Y  %H:%M')}  •  Confidential"
    )
    foot.alignment          = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size  = Pt(8.5)
    foot.runs[0].italic     = True
    foot.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Excel builder ──────────────────────────────────────────────────────────────

def generate_excel_report(tickets_data: list) -> io.BytesIO:
    """Build a styled .xlsx audit log and return it as a BytesIO buffer."""
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    buf    = io.BytesIO()
    df_log = pd.DataFrame(tickets_data)

    # Add human-readable response time column next to the raw seconds
    if "Response_Time_Secs" in df_log.columns:
        def _fmt_rt(secs):
            try:
                s = int(secs or 0)
                return f"{s // 60}m {s % 60}s" if s else "—"
            except (ValueError, TypeError):
                return "—"
        idx = df_log.columns.tolist().index("Response_Time_Secs") + 1
        df_log.insert(idx, "Response_Time", df_log["Response_Time_Secs"].apply(_fmt_rt))

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:

        # Sheet 1: Audit Log
        df_log.to_excel(writer, sheet_name="Audit Log", index=False)
        ws = writer.sheets["Audit Log"]

        header_fill = PatternFill("solid", fgColor="1F4E79")
        header_font = Font(bold=True, color="FFFFFF", size=10)
        for cell in ws[1]:
            cell.fill      = header_fill
            cell.font      = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

        sev_fills = {
            "CRITICAL": PatternFill("solid", fgColor="FF0000"),
            "HIGH":     PatternFill("solid", fgColor="FF8C00"),
            "MEDIUM":   PatternFill("solid", fgColor="FFC000"),
            "LOW":      PatternFill("solid", fgColor="70AD47"),
        }
        white_font   = Font(bold=True, color="FFFFFF", size=10)
        headers      = [c.value for c in ws[1]]
        sev_col_idx  = headers.index("Severity") + 1 if "Severity" in headers else None
        thin         = Side(style="thin", color="D9D9D9")
        border       = Border(left=thin, right=thin, top=thin, bottom=thin)

        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            for cell in row:
                cell.alignment = Alignment(vertical="center")
                cell.border    = border
                if sev_col_idx and cell.column == sev_col_idx:
                    sev_val = str(cell.value or "").upper()
                    fill    = sev_fills.get(sev_val)
                    if fill:
                        cell.fill = fill
                        cell.font = white_font

        for col_idx, col_cells in enumerate(ws.columns, 1):
            max_len = max((len(str(c.value or "")) for c in col_cells), default=8)
            ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 40)
        ws.row_dimensions[1].height = 20
        ws.freeze_panes = "A2"

        # Sheet 2: Summary
        total         = len(tickets_data)
        approved      = sum(1 for t in tickets_data if t["Status"] in ("Approved", "Approved (Queue)"))
        dropped       = sum(1 for t in tickets_data if t["Status"] == "Dropped (Duplicate)")
        rejected      = sum(1 for t in tickets_data if t["Status"] == "Rejected")
        sla_breached  = sum(1 for t in tickets_data if t.get("SLA_Breached", False))
        critical_cnt  = sum(1 for t in tickets_data if str(t.get("Severity","")).upper() == "CRITICAL")
        high_cnt      = sum(1 for t in tickets_data if str(t.get("Severity","")).upper() == "HIGH")
        rt_vals       = [int(t.get("Response_Time_Secs", 0) or 0) for t in tickets_data]
        avg_rt        = int(sum(rt_vals) / len(rt_vals)) if rt_vals else 0
        avg_rt_str    = f"{avg_rt // 60}m {avg_rt % 60}s" if avg_rt else "—"
        comp_pct      = f"{round((1 - sla_breached / total) * 100)}%" if total else "—"
        conf_vals     = [float(t["Confidence_Score"]) for t in tickets_data
                         if str(t.get("Confidence_Score","")).strip().lstrip("-").isdigit()]
        avg_conf      = f"{round(sum(conf_vals)/len(conf_vals))}%" if conf_vals else "—"
        cat_counts: dict = {}
        for t in tickets_data:
            c = t.get("Category", "Unknown")
            cat_counts[c] = cat_counts.get(c, 0) + 1

        summary_rows = [
            ("Metric", "Value"),
            ("Total Processed",         total),
            ("Approved & Escalated",    approved),
            ("Duplicates Dropped",      dropped),
            ("Rejected",                rejected),
            ("SLA Breached",            sla_breached),
            ("SLA Compliance Rate",     comp_pct),
            ("Critical Incidents",      critical_cnt),
            ("High-Severity",           high_cnt),
            ("Avg. Response Time",      avg_rt_str),
            ("Avg. AI Confidence",      avg_conf),
            ("", ""),
            ("Category Breakdown",      ""),
        ] + [(cat, cnt) for cat, cnt in sorted(cat_counts.items())]

        df_summary = pd.DataFrame(summary_rows[1:], columns=summary_rows[0])
        df_summary.to_excel(writer, sheet_name="Summary", index=False)
        ws2 = writer.sheets["Summary"]
        for cell in ws2[1]:
            cell.fill      = header_fill
            cell.font      = header_font
            cell.alignment = Alignment(horizontal="center")
        for col_cells in ws2.columns:
            max_len = max((len(str(c.value or "")) for c in col_cells), default=8)
            ws2.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 35)

        # Sheet 3: By Category
        if not df_log.empty and "Category" in df_log.columns:
            cat_df = df_log.groupby(["Category", "Status"]).size().unstack(fill_value=0).reset_index()
            cat_df.to_excel(writer, sheet_name="By Category", index=False)
            ws3 = writer.sheets["By Category"]
            for cell in ws3[1]:
                cell.fill = header_fill
                cell.font = header_font
            for col_cells in ws3.columns:
                max_len = max((len(str(c.value or "")) for c in col_cells), default=8)
                ws3.column_dimensions[col_cells[0].column_letter].width = min(max_len + 4, 25)

    buf.seek(0)
    return buf
